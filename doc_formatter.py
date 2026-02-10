#!/usr/bin/env python3
"""
DocGen - Document Formatter

Format Word documents according to defined styles (fonts, sizes, spacing, etc.)
Support both DOCX input and text/markdown input.
"""

import json
import os
import re
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentFormatter:
    """Format documents according to defined style rules."""
    
    def __init__(self, style_config: Dict[str, Any] = None):
        """Initialize with style configuration."""
        self.config = style_config or self.get_default_style()
    
    def get_default_style(self) -> Dict[str, Any]:
        """Get default Chinese document style (GB/T 9704-2012)."""
        return {
            "document": {
                "margin_top": Cm(3.7),
                "margin_bottom": Cm(3.5),
                "margin_left": Cm(2.8),
                "margin_right": Cm(2.6),
                "line_spacing": 1.5,
                "font_family": "FangSong_GB2312",
                "font_size": 16
            },
            "title": {
                "font_family": "SimHei",
                "font_size": 22,
                "bold": True,
                "alignment": "center",
                "spacing_before": 400,
                "spacing_after": 20
            },
            "heading1": {
                "font_family": "SimHei",
                "font_size": 16,
                "bold": True,
                "alignment": "left",
                "spacing_before": 200,
                "spacing_after": 10
            },
            "heading2": {
                "font_family": "KaiTi_GB2312",
                "font_size": 15,
                "bold": False,
                "alignment": "left",
                "spacing_before": 100,
                "spacing_after": 8
            },
            "body": {
                "font_family": "FangSong_GB2312",
                "font_size": 16,
                "bold": False,
                "alignment": "left",
                "first_line_indent": 2,
                "line_spacing": 1.5,
                "spacing_before": 0,
                "spacing_after": 0
            },
            "signature": {
                "font_family": "FangSong_GB2312",
                "font_size": 16,
                "bold": False,
                "alignment": "right",
                "spacing_before": 200,
                "spacing_after": 0
            }
        }
    
    def apply_style_to_paragraph(self, para, style_config: Dict[str, Any]):
        """Apply style to a paragraph."""
        if not para.runs:
            return
        
        # Font family
        if 'font_family' in style_config:
            para.runs[0].font.name = style_config['font_family']
            rPr = para.runs[0]._r.get_or_add_rPr()
            rPr.get_or_add_rFonts().east_asia = style_config['font_family']
        
        # Font size
        if 'font_size' in style_config:
            para.runs[0].font.size = Pt(style_config['font_size'])
        
        # Bold
        if 'bold' in style_config:
            para.runs[0].font.bold = style_config['bold']
        
        # Alignment
        if 'alignment' in style_config:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if style_config['alignment'] in align_map:
                para.alignment = align_map[style_config['alignment']]
        
        # Spacing
        if 'spacing_before' in style_config:
            para.paragraph_format.space_before = Pt(style_config['spacing_before'] / 20)
        if 'spacing_after' in style_config:
            para.paragraph_format.space_after = Pt(style_config['spacing_after'] / 20)
        
        # Line spacing
        if 'line_spacing' in style_config:
            para.paragraph_format.line_spacing = style_config['line_spacing']
        
        # First line indent
        if 'first_line_indent' in style_config:
            para.paragraph_format.first_line_indent = Cm(style_config['first_line_indent'] * 0.5)
    
    def detect_paragraph_type(self, para: Document.paragraphs) -> str:
        """Detect the type of paragraph based on its style and content."""
        # Check style name
        style_name = para.style.name if para.style else ""
        if 'Heading' in style_name or 'Title' in style_name:
            return 'heading1'
        
        # Check by font size (rough estimation)
        if para.runs and para.runs[0].font.size:
            size = para.runs[0].font.size.pt
            if size >= 22:
                return 'title'
            elif size >= 18:
                return 'heading1'
            elif size >= 16:
                return 'heading2'
            else:
                return 'body'
        
        # Check by content patterns
        text = para.text.strip()
        if not text:
            return 'empty'
        
        # Chinese numbering patterns: 一、二、三... or 1、2、3...
        if re.match(r'^[一二三四五六七八九十\d]+\s*[\.\、：:]', text):
            return 'heading1'
        if re.match(r'^[（\(]\s*\d+[）\)]', text):
            return 'heading2'
        
        return 'body'
    
    def format_word_document(self, input_path: str, output_path: str) -> str:
        """Read a Word document, detect styles, and reformat."""
        # Read source document
        source_doc = Document(input_path)
        
        # Create new document
        new_doc = Document()
        
        # Apply document-level settings
        doc_config = self.config.get('document', {})
        section = new_doc.sections[0]
        section.top_margin = doc_config.get('margin_top', Cm(3.7))
        section.bottom_margin = doc_config.get('margin_bottom', Cm(3.5))
        section.left_margin = doc_config.get('margin_left', Cm(2.8))
        section.right_margin = doc_config.get('margin_right', Cm(2.6))
        
        # Process each paragraph
        for para in source_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect paragraph type
            para_type = self.detect_paragraph_type(para)
            
            # Get config
            style_map = {
                'title': self.config.get('title', {}),
                'heading1': self.config.get('heading1', {}),
                'heading2': self.config.get('heading2', {}),
                'body': self.config.get('body', {}),
            }
            
            style_config = style_map.get(para_type, self.config.get('body', {}))
            
            # Create new paragraph with same text
            new_para = new_doc.add_paragraph(text)
            
            # Apply style
            self.apply_style_to_paragraph(new_para, style_config)
        
        # Save
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        new_doc.save(str(output))
        
        return str(output_path)
    
    def format_document(self, content: str, output_path: str) -> str:
        """Format content from text/markdown."""
        doc = Document()
        
        doc_config = self.config.get('document', {})
        section = doc.sections[0]
        section.top_margin = doc_config.get('margin_top', Cm(3.7))
        section.bottom_margin = doc_config.get('margin_bottom', Cm(3.5))
        section.left_margin = doc_config.get('margin_left', Cm(2.8))
        section.right_margin = doc_config.get('margin_right', Cm(2.6))
        
        lines = content.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect line type by markdown markers
            if line.startswith('# '):
                para = doc.add_heading(line[2:].strip(), 0)
                self.apply_style_to_paragraph(para, self.config.get('title', {}))
            elif line.startswith('## '):
                para = doc.add_paragraph(line[3:].strip())
                self.apply_style_to_paragraph(para, self.config.get('heading1', {}))
            elif line.startswith('### '):
                para = doc.add_paragraph(line[4:].strip())
                self.apply_style_to_paragraph(para, self.config.get('heading2', {}))
            elif line.startswith('---'):
                para = doc.add_paragraph(line[3:].strip() if len(line) > 3 else '')
                self.apply_style_to_paragraph(para, self.config.get('signature', {}))
            else:
                para = doc.add_paragraph(line)
                self.apply_style_to_paragraph(para, self.config.get('body', {}))
        
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        
        return str(output_path)
    
    def format_from_file(self, input_file: str, output_path: str) -> str:
        """Read content from file and format it."""
        ext = Path(input_file).suffix.lower()
        
        if ext == '.docx':
            return self.format_word_document(input_file, output_path)
        else:
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            return self.format_document(content, output_path)


def main():
    """CLI interface."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='DocGen - Format Word documents with professional styles'
    )
    parser.add_argument('input', nargs='?', help='Input file (.docx, .md, .txt)')
    parser.add_argument('-o', '--output', default='output.docx', help='Output filename')
    parser.add_argument('-c', '--config', help='Style config JSON file')
    parser.add_argument('-l', '--list', action='store_true', help='Show available styles')
    parser.add_argument('--preview', action='store_true', help='Preview style settings')
    
    args = parser.parse_args()
    
    if args.list:
        print("Available styles:")
        print("  - default  : GB/T 9704-2012 Chinese document standard")
        print("  - formal   : Formal business document")
        return
    
    if args.preview:
        formatter = DocumentFormatter()
        print("Current style configuration:")
        print(json.dumps(formatter.config, ensure_ascii=False, indent=2))
        return
    
    if not args.input:
        parser.print_help()
        return
    
    # Load custom config if provided
    config = None
    if args.config and os.path.exists(args.config):
        with open(args.config, 'r', encoding='utf-8') as f:
            config = json.load(f)
    
    formatter = DocumentFormatter(config)
    
    try:
        output_path = formatter.format_from_file(args.input, args.output)
        print("[DocGen] Document formatted successfully!")
        print(f"  Input:  {args.input}")
        print(f"  Output: {output_path}")
    except Exception as e:
        print(f"Error: {e}")


if __name__ == '__main__':
    main()
