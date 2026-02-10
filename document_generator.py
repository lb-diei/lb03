#!/usr/bin/env python3
"""
Document Template System

Generate professional DOCX documents from customizable templates.
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class DocumentGenerator:
    """Generate DOCX documents from templates."""
    
    def __init__(self, template_dir: str = "templates"):
        """Initialize with template directory."""
        self.template_dir = Path(template_dir)
        self.default_vars = {
            "title": "Document Title",
            "author": "Author Name",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "content": "Your content here...",
        }
    
    def load_template(self, template_name: str) -> Dict[str, str]:
        """Load a template file."""
        template_path = self.template_dir / f"{template_name}.md"
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        with open(template_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse simple key-value format
        variables = {}
        lines = content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if line.startswith('[') and line.endswith(']'):
                current_section = line[1:-1]
            elif ':' in line and current_section == "variables":
                key, value = line.split(':', 1)
                variables[key.strip()] = value.strip()
            elif line.startswith('# '):
                # Title
                variables['title'] = line[2:]
            elif current_section == "content":
                if 'content' not in variables:
                    variables['content'] = ""
                variables['content'] += line + "\n"
        
        return variables
    
    def generate_document(
        self,
        template_name: str,
        output_name: str,
        variables: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Generate a DOCX document from a template."""
        # Load template variables
        template_vars = self.load_template(template_name)
        
        # Merge with provided variables (user input overrides template defaults)
        if variables:
            template_vars.update(variables)
        
        # Use default for missing values
        for key, value in self.default_vars.items():
            if key not in template_vars:
                template_vars[key] = value
        
        # Create document
        doc = Document()
        
        # Set title
        title = doc.add_heading(template_vars.get('title', 'Document'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if 'author' in template_vars:
            author_para = doc.add_paragraph()
            author_para.add_run(f"Author: {template_vars['author']}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if 'date' in template_vars:
            date_para = doc.add_paragraph()
            date_para.add_run(f"Date: {template_vars['date']}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add separator
        doc.add_paragraph("_" * 50)
        
        # Add content
        content = template_vars.get('content', '')
        if content:
            # Split by double newlines for paragraphs
            paragraphs = content.strip().split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        
        # Save document
        output_path = Path(output_name)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return str(output_path)
    
    def list_templates(self) -> list:
        """List available templates."""
        if not self.template_dir.exists():
            return []
        
        templates = []
        for f in self.template_dir.glob('*.md'):
            templates.append(f.stem)
        
        return templates


def main():
    """CLI interface."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Document Template System')
    parser.add_argument('template', nargs='?', help='Template name (without .md extension)')
    parser.add_argument('-o', '--output', default='output.docx', help='Output filename')
    parser.add_argument('-l', '--list', action='store_true', help='List available templates')
    parser.add_argument('-v', '--variable', action='append', help='Variable in format key=value')
    
    args = parser.parse_args()
    
    generator = DocumentGenerator()
    
    if args.list:
        templates = generator.list_templates()
        if templates:
            print("Available templates:")
            for t in templates:
                print(f"  - {t}")
        else:
            print("No templates found in 'templates/' directory.")
        return
    
    if not args.template:
        parser.print_help()
        return
    
    # Parse variables
    variables = {}
    if args.variable:
        for var in args.variable:
            if '=' in var:
                key, value = var.split('=', 1)
                variables[key] = value
    
    # Generate document
    try:
        output_path = generator.generate_document(
            args.template,
            args.output,
            variables,
        )
        print(f"Document generated: {output_path}")
    except Exception as e:
        print(f"Error: {e}")


if __name__ == '__main__':
    main()
